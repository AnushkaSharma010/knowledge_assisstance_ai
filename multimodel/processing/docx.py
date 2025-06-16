import docx
from PIL import Image
import io
import pandas as pd
import os
import zipfile
import xml.etree.ElementTree as ET
from typing import List, Dict, Union
from core.captioner import GeminiMultimodalProcessor
from schemas import DocumentChunk
from logger import get_logger

logger = get_logger("DocxProcessor")

class DocxProcessor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = docx.Document(file_path)
        self.captioner = GeminiMultimodalProcessor()

    def process(self) -> List[DocumentChunk]:
        logger.info(f"Processing DOCX file: {self.file_path}")
        chunks = []

        text_parts = []

        # Extract paragraphs
        for para in self.document.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extract tables
        for i, table in enumerate(self.document.tables):
            try:
                data = [[cell.text for cell in row.cells] for row in table.rows]
                df = pd.DataFrame(data[1:], columns=data[0]) if data else pd.DataFrame()
                markdown = df.to_markdown(index=False)
                logger.debug(f"Extracted table {i+1} with shape {df.shape}")
                text_parts.append(f"\nTable:\n{markdown}")
            except Exception as e:
                logger.warning(f"Failed to process table {i+1}: {e}")

        # Extract images from zip structure
        try:
            with zipfile.ZipFile(self.file_path) as z:
                with z.open('word/document.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    ns = {
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                    }
                    blips = root.findall('.//a:blip', ns)

                with z.open('word/_rels/document.xml.rels') as rels_file:
                    rels_tree = ET.parse(rels_file)
                    rels_root = rels_tree.getroot()
                    nsmap = {'pr': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                    rels = {rel.attrib['Id']: rel.attrib['Target'] for rel in rels_root.findall('pr:Relationship', nsmap)}

                for i, elem in enumerate(blips):
                    r_id = elem.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if not r_id or r_id not in rels:
                        continue
                    image_path = f'word/{rels[r_id]}'
                    with z.open(image_path) as img_file:
                        img_data = img_file.read()
                        try:
                            caption_obj = self.captioner.process(img_data, 'image')
                            caption = getattr(caption_obj, 'caption', '')
                            logger.debug(f"Caption for image {i+1}: {caption}")
                            text_parts.append(f"\nImage: {caption}")
                        except Exception as e:
                            logger.warning(f"Failed to process image {i+1}: {e}")
        except Exception as e:
            logger.warning(f"Failed to extract images from DOCX: {e}")

        combined = "\n\n".join(text_parts)
        if combined.strip():
            chunks.append(DocumentChunk(
                content=combined.strip(),
                type="text",
                page_number=None
            ))
            logger.info(f"Created 1 unified chunk with length {len(combined.strip())}")
        else:
            logger.warning("No content extracted from DOCX file.")

        return chunks